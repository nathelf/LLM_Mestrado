# ✅ Documentação Técnica de Agentes Valorant com Melhorias Profissionais

# Instalação dos pacotes (Colab)
!pip install -q transformers accelerate bitsandbytes markdown2 python-docx

from transformers import AutoTokenizer, AutoModelForCausalLM, BitsAndBytesConfig, pipeline
import torch
import os
import json
from docx import Document
from google.colab import files
from zipfile import ZipFile
import shutil
import gc
from datetime import datetime

# Diretórios
os.makedirs("json_valorant", exist_ok=True)
os.makedirs("resultados", exist_ok=True)
os.makedirs("docx", exist_ok=True)

# Upload de arquivos JSON
print("🔼 Envie os arquivos .json dos agentes (um por arquivo)...")
uploaded = files.upload()
for nome in uploaded:
    shutil.move(nome, f"json_valorant/{nome}")

# Prompt detalhado
prompt_base = '''Sua tarefa é analisar um objeto JSON contendo dados de um agente da API pública do jogo Valorant e gerar uma documentação técnica clara, precisa e adequada para desenvolvedores.

---

1. **Título:** Nome do agente (campo `displayName`)
2. **Descrição geral** explicando o papel e o uso do agente no jogo.
3. **Tabela com os principais campos**:
   | Campo | Tipo de dado | Descrição |
   |-------|--------------|-----------|
   Preencha com os campos principais do JSON.
4. **Subseções** para objetos aninhados como `role`, `abilities` e outros.
5. Use Markdown estruturado (títulos `##`, listas `-` e tabelas).
6. Não repita o JSON nem inclua comentários irrelevantes.

JSON de entrada:

{{json}}

Gere a documentação de forma clara, objetiva e técnica.
'''

# Modelo e pipeline com quantização e offload
model_id = "HuggingFaceH4/zephyr-7b-beta"
tokenizer = AutoTokenizer.from_pretrained(model_id)
bnb_config = BitsAndBytesConfig(
    load_in_8bit=True,
    llm_int8_enable_fp32_cpu_offload=True
)
model = AutoModelForCausalLM.from_pretrained(
    model_id,
    device_map="auto",  # ou "cpu" se estiver sem GPU
    torch_dtype=torch.float16  # ou .float32 se quiser compatibilidade máxima
)

pipe = pipeline("text-generation", model=model, tokenizer=tokenizer, max_new_tokens=1024)

# Processamento
for filename in os.listdir("json_valorant"):
    if filename.endswith(".json"):
        with open(f"json_valorant/{filename}", "r", encoding="utf-8") as f:
            data = json.load(f)

        prompt = prompt_base.replace("{{json}}", json.dumps(data, indent=2, ensure_ascii=False))
        result = pipe(prompt)[0]["generated_text"]

        # Markdown intermediário
        md_path = f"resultados/{filename.replace('.json', '.md')}"
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(result)

        # Documento Word
        doc = Document()
        doc.add_heading(f"Documentação Técnica – {data.get('displayName', filename)}", level=1)
        doc.add_paragraph("Projeto: API Valorant", style="Intense Quote")
        doc.add_paragraph("Data: " + datetime.now().strftime("%d/%m/%Y"), style="Intense Quote")
        doc.add_paragraph("\n")

        # Processa cada linha do markdown
        for linha in result.split("\n"):
            linha = linha.strip()
            if linha.startswith("##"):
                doc.add_heading(linha.replace("##", "").strip(), level=2)
            elif linha.startswith("#"):
                doc.add_heading(linha.replace("#", "").strip(), level=3)
            elif linha.startswith("|") and "|" in linha[1:]:
                # Cria tabela a partir do markdown
                if "Campo" in linha:
                    colunas = [c.strip() for c in linha.split("|") if c.strip()]
                    tabela = doc.add_table(rows=1, cols=len(colunas))
                    tabela.style = 'Table Grid'
                    for i, coluna in enumerate(colunas):
                        tabela.cell(0, i).text = coluna
                elif "---" not in linha:
                    valores = [c.strip() for c in linha.split("|") if c.strip()]
                    row = tabela.add_row()
                    for i, val in enumerate(valores):
                        row.cells[i].text = val
            elif linha.startswith("-") or linha.startswith("*"):
                doc.add_paragraph(linha[1:].strip(), style="List Bullet")
            elif linha:
                doc.add_paragraph(linha)

        docx_path = f"docx/{filename.replace('.json', '.docx')}"
        doc.save(docx_path)

        gc.collect()
        torch.cuda.empty_cache()

# Compacta documentos gerados
with ZipFile("documentacoes_valorant.zip", "w") as zipf:
    for file in os.listdir("docx"):
        zipf.write(f"docx/{file}", arcname=file)

files.download("documentacoes_valorant.zip")
